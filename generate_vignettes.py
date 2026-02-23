# -*- coding: utf-8 -*-
#!/usr/bin/env python3
"""
MedScrape Vignette Generator
-----------------------------
Generates one-page publication-ready vignettes for 5 nutrition-focused PubMed
searches, demonstrating the efficacy of the MedScrape AI-assisted review method.

Usage:
    python generate_vignettes.py your.email@institution.edu
    python generate_vignettes.py your.email@institution.edu --max 3000
    python generate_vignettes.py your.email@institution.edu --start 2015 --end 2025

Outputs:
    vignettes/vignette_search1_AI_and_Nutrition.pdf
    vignettes/vignette_search2_Precision_Nutrition_and_Genetics.pdf
    ... (one PDF per search + one CSV per search in vignettes/data/)

Authors: Department of Mathematical Sciences, USMA West Point
"""

import sys
import os
import io
import argparse
import textwrap
import time
import warnings
from datetime import datetime

warnings.filterwarnings('ignore')

# -- Ensure the medrev directory is on the path -------------------------------
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import matplotlib.patches as mpatches
from matplotlib.colors import to_rgba

from Bio import Entrez

# NCBI API key — raises rate limit from 3 req/s to 10 req/s
NCBI_API_KEY = "5a3c321d881ae6f0902a1ba96538a24e4808"

try:
    from clustering_enhanced import EnhancedClusterer, ReviewComparator
    CLUSTERING_OK = True
except ImportError as _e:
    CLUSTERING_OK = False
    print(f"[WARN] clustering_enhanced.py not found ({_e}). "
          "Vignettes will be generated without cluster analysis.")


# -----------------------------------------------------------------------------
#  Search Definitions
# -----------------------------------------------------------------------------

SEARCHES = [
    {
        "id": 1,
        "title": "AI & Nutrition",
        "short": "AI_and_Nutrition",
        "query": '("Artificial Intelligence" OR "AI") AND "Nutrition"',
        "description": (
            "Explores artificial intelligence applications across nutrition science, "
            "including dietary assessment, personalized recommendations, and metabolic modeling."
        ),
    },
    {
        "id": 2,
        "title": "Precision Nutrition & Genetics",
        "short": "Precision_Nutrition_Genetics",
        "query": '"Precision Nutrition" AND ("Genetics" AND "Nutrition")',
        "description": (
            "Investigates genotype-informed dietary strategies, nutrigenomics, "
            "and SNP-based personalized nutrition interventions."
        ),
    },
    {
        "id": 3,
        "title": "Precision Nutrition & Exercise",
        "short": "Precision_Nutrition_Exercise",
        "query": '"Precision Nutrition" AND ("Exercise" AND "Nutrition")',
        "description": (
            "Examines individualized nutrition-exercise interactions, sport-specific "
            "dietary prescriptions, and performance optimization through personalized fueling."
        ),
    },
    {
        "id": 4,
        "title": "Precision Nutrition & Microbiome",
        "short": "Precision_Nutrition_Microbiome",
        "query": '"Precision Nutrition" AND ("Microbiome" AND "Nutrition")',
        "description": (
            "Covers gut microbiota-diet relationships, microbiome-guided dietary "
            "interventions, and host-microbiome interactions in precision health."
        ),
    },
    {
        "id": 5,
        "title": "Dietary Supplements & Diabetes",
        "short": "Dietary_Supplements_Diabetes",
        "query": '"Dietary Supplements" AND "Diabetes"',
        "description": (
            "Surveys supplementation strategies for glycemic control, insulin sensitivity, "
            "and diabetes prevention including vitamins, minerals, and botanical compounds."
        ),
    },
]

# Colour palette (consistent across all vignettes)
PALETTE = {
    "header_bg":   "#1a3a5c",
    "header_text": "#ffffff",
    "subtext":     "#c8ddf0",
    "accent":      "#2980b9",
    "gap_red":     "#c0392b",
    "gap_green":   "#27ae60",
    "bar":         "#5b9bd5",
    "trend":       "#e74c3c",
    "type_colors": ["#2196F3", "#4CAF50", "#FF9800", "#9C27B0", "#F44336", "#00BCD4"],
}


# -----------------------------------------------------------------------------
#  PubMed Scraping  (mirrors app_enhanced.py, CLI-friendly)
# -----------------------------------------------------------------------------

def _scrape(query: str, email: str, max_results: int,
            start_year: int, end_year: int,
            pub_type: str = "all") -> pd.DataFrame:
    """Query PubMed and return a DataFrame.  pub_type: 'all' | 'review' | 'non-review'"""
    Entrez.email = email
    Entrez.api_key = NCBI_API_KEY

    full_q = query
    if pub_type == "review":
        full_q += ' AND "review"[Publication Type]'
    elif pub_type == "non-review":
        full_q += ' NOT "review"[Publication Type]'
    full_q += f" AND {start_year}:{end_year}[pdat]"

    # usehistory stores the result set server-side — no ID-list cap
    try:
        handle = Entrez.esearch(db="pubmed", term=full_q, usehistory="y")
        rec = Entrez.read(handle)
        total_count = int(rec["Count"])
        webenv    = rec["WebEnv"]
        query_key = rec["QueryKey"]
        handle.close()
    except Exception as e:
        print(f"    [ERROR] esearch failed: {e}")
        return pd.DataFrame()

    fetch_count = min(total_count, max_results)
    if fetch_count == 0:
        print("    No results found.")
        return pd.DataFrame()

    print(f"    PubMed total matches: {total_count:,}  |  Fetching: {fetch_count:,}")

    rows = []
    batch = 500
    for retstart in range(0, fetch_count, batch):
        batch_end = min(retstart + batch, fetch_count)
        print(f"    Records {retstart + 1}-{batch_end} / {fetch_count:,} ...", end="\r", flush=True)
        try:
            h2 = Entrez.efetch(
                db="pubmed",
                rettype="medline", retmode="xml",
                retstart=retstart, retmax=batch,
                webenv=webenv, query_key=query_key,
            )
            recs = Entrez.read(h2)
            h2.close()
        except Exception as e:
            print(f"\n    [WARN] Batch fetch error: {e}")
            time.sleep(2)
            continue

        for r in recs.get("PubmedArticle", []):
            try:
                art = r["MedlineCitation"]["Article"]

                authors = []
                for a in art.get("AuthorList", []):
                    if "LastName" in a and "Initials" in a:
                        authors.append(f"{a['LastName']} {a['Initials']}")

                abstract = ""
                if "Abstract" in art:
                    parts = art["Abstract"].get("AbstractText", [])
                    abstract = " ".join(str(p) for p in parts)

                year = None
                if art.get("ArticleDate"):
                    year = int(art["ArticleDate"][0]["Year"])
                elif "JournalIssue" in art.get("Journal", {}):
                    pd_ = art["Journal"]["JournalIssue"].get("PubDate", {})
                    if "Year" in pd_:
                        year = int(pd_["Year"])

                tl = art.get("PublicationTypeList", [])
                longitudinal_kw = ["longitudinal", "prospective cohort", "follow-up", "follow up"]

                rows.append({
                    "PMID":            str(r["MedlineCitation"]["PMID"]),
                    "Title":           art.get("ArticleTitle", ""),
                    "Abstract":        abstract,
                    "Authors":         "; ".join(authors),
                    "Journal":         art.get("Journal", {}).get("Title", ""),
                    "Year":            year,
                    "URL":             f"https://pubmed.ncbi.nlm.nih.gov/{r['MedlineCitation']['PMID']}/",
                    "Review":          int(any(t.lower() == "review" for t in tl)),
                    "SystematicReview": int(any(t.lower() == "systematic review" for t in tl)),
                    "ClinicalTrial":   int(any(t.lower() == "clinical trial" for t in tl)),
                    "MetaAnalysis":    int(any(t.lower() == "meta-analysis" for t in tl)),
                    "RCT":             int(any(t.lower() == "randomized controlled trial" for t in tl)),
                    "LongitudinalStudy": int(any(kw in abstract.lower() for kw in longitudinal_kw)),
                })
            except Exception:
                continue

        time.sleep(0.11)  # ~9 req/s — safely under the 10 req/s API-key limit

    print()  # newline after \r
    return pd.DataFrame(rows)


# -----------------------------------------------------------------------------
#  Vignette Drawing
# -----------------------------------------------------------------------------

def _wrap(text: str, width: int = 80) -> str:
    return "\n".join(textwrap.wrap(text, width))


def draw_vignette(search: dict,
                  df_all: pd.DataFrame,
                  review_results=None,
                  non_review_results=None,
                  similarity_matrix=None,
                  gap_indices=None,
                  output_path: str = None) -> str:
    """
    Render a single one-page vignette as a PDF / PNG.

    Layout (8.5 x 11 in, letter):
    +--------------------------------------------------------+
    |  HEADER: title | query | key stats                     |
    |----------------------+---------------------------------|
    |  Publication trend   |  Article-type distribution      |
    |----------------------+---------------------------------|
    |  Cluster scatter     |  Research themes (top terms)    |
    |----------------------+---------------------------------|
    |  Gap-analysis banner                                   |
    +--------------------------------------------------------+
    """
    fig = plt.figure(figsize=(8.5, 11), dpi=150, facecolor="white")

    gs = gridspec.GridSpec(
        4, 2, figure=fig,
        height_ratios=[0.85, 2.1, 3.0, 1.6],
        hspace=0.50, wspace=0.32,
        left=0.07, right=0.96, top=0.97, bottom=0.04,
    )

    # -- 1  HEADER -------------------------------------------------------------
    ax_h = fig.add_subplot(gs[0, :])
    ax_h.set_axis_off()

    ax_h.add_patch(mpatches.FancyBboxPatch(
        (0, 0), 1, 1,
        boxstyle="round,pad=0.02",
        lw=0, facecolor=PALETTE["header_bg"],
        transform=ax_h.transAxes, zorder=0,
    ))

    ax_h.text(0.5, 0.88,
              f"Search {search['id']}:  {search['title']}",
              transform=ax_h.transAxes, ha="center", va="top",
              fontsize=12, fontweight="bold", color=PALETTE["header_text"], zorder=1)

    query_wrapped = _wrap(f"PubMed Query:  {search['query']}", 90)
    ax_h.text(0.5, 0.58, query_wrapped,
              transform=ax_h.transAxes, ha="center", va="top",
              fontsize=7, color=PALETTE["subtext"], style="italic", zorder=1)

    df_c = df_all[df_all["Year"].notna()].copy()
    yr_min = int(df_c["Year"].min()) if not df_c.empty else "--"
    yr_max = int(df_c["Year"].max()) if not df_c.empty else "--"
    n_rev  = int(df_all.get("Review", pd.Series(dtype=int)).sum())
    stats  = (f"n = {len(df_all):,} articles   |   "
              f"Years: {yr_min}-{yr_max}   |   "
              f"Journals: {df_all['Journal'].nunique():,}   |   "
              f"Reviews: {n_rev:,}")
    ax_h.text(0.5, 0.10, stats,
              transform=ax_h.transAxes, ha="center", va="bottom",
              fontsize=7.5, fontweight="bold", color=PALETTE["subtext"], zorder=1)

    # -- 2  PUBLICATION TREND (left) -------------------------------------------
    ax_t = fig.add_subplot(gs[1, 0])

    if not df_c.empty and df_c["Year"].nunique() >= 2:
        yc = df_c["Year"].value_counts().sort_index()
        years  = yc.index.astype(int).values
        counts = yc.values

        ax_t.bar(years, counts, color=PALETTE["bar"], alpha=0.80, width=0.85, zorder=2)

        if len(years) >= 3:
            z = np.polyfit(years, counts, 1)
            xs = np.linspace(years[0], years[-1], 200)
            ax_t.plot(xs, np.poly1d(z)(xs), "--",
                      color=PALETTE["trend"], lw=1.5, label="Trend", zorder=3)
            ax_t.legend(fontsize=6.5, loc="upper left", framealpha=0.7)

        # Growth annotation
        if len(counts) >= 4:
            thirds = max(1, len(counts) // 3)
            early  = counts[:thirds].mean()
            late   = counts[-thirds:].mean()
            if early > 0:
                pct  = (late - early) / early * 100
                sym  = "^" if pct > 0 else "v"
                col  = PALETTE["gap_green"] if pct > 0 else PALETTE["gap_red"]
                ax_t.text(0.97, 0.95, f"{sym} {abs(pct):.0f}% growth",
                          transform=ax_t.transAxes, ha="right", va="top",
                          fontsize=7, color=col,
                          bbox=dict(boxstyle="round,pad=0.25", fc="white",
                                    ec="lightgray", alpha=0.85))

        tick_step = max(1, (years[-1] - years[0]) // 6)
        tick_yrs  = list(range(years[0], years[-1] + 1, tick_step))
        if years[-1] not in tick_yrs:
            tick_yrs.append(int(years[-1]))
        ax_t.set_xticks(tick_yrs)
        ax_t.set_xticklabels(tick_yrs, rotation=45, ha="right", fontsize=6.5)
    else:
        ax_t.text(0.5, 0.5, "Insufficient year data",
                  transform=ax_t.transAxes, ha="center", va="center",
                  fontsize=9, color="gray")

    ax_t.set_ylabel("Publications / year", fontsize=7.5)
    ax_t.set_title("Publication Trend", fontsize=9, fontweight="bold", pad=5)
    ax_t.tick_params(axis="y", labelsize=6.5)
    ax_t.spines[["top", "right"]].set_visible(False)

    # -- 3  ARTICLE-TYPE DISTRIBUTION (right) ---------------------------------
    ax_ty = fig.add_subplot(gs[1, 1])

    type_labels = ["Reviews", "Syst. Reviews", "Clin. Trials",
                   "Meta-Analyses", "RCTs", "Longitudinal"]
    type_cols   = ["Review", "SystematicReview", "ClinicalTrial",
                   "MetaAnalysis", "RCT", "LongitudinalStudy"]
    counts_ty   = [int(df_all[c].sum()) if c in df_all.columns else 0 for c in type_cols]

    h_bars = ax_ty.barh(type_labels, counts_ty,
                        color=PALETTE["type_colors"], edgecolor="white", height=0.62)
    for bar, cnt in zip(h_bars, counts_ty):
        if cnt > 0:
            ax_ty.text(bar.get_width() + max(counts_ty, default=1) * 0.03,
                       bar.get_y() + bar.get_height() / 2,
                       f"{cnt:,}", va="center", ha="left",
                       fontsize=6.5, color="#333")

    ax_ty.set_xlabel("Count", fontsize=7.5)
    ax_ty.set_title("Article-Type Distribution", fontsize=9, fontweight="bold", pad=5)
    ax_ty.tick_params(axis="y", labelsize=7)
    ax_ty.tick_params(axis="x", labelsize=6.5)
    ax_ty.spines[["top", "right"]].set_visible(False)
    ax_ty.set_xlim(0, max(counts_ty, default=1) * 1.20)

    # -- 4  CLUSTER SCATTER PLOT (left) ---------------------------------------
    ax_sc = fig.add_subplot(gs[2, 0])
    ax_sc.set_title("Research Cluster Map (PCA)", fontsize=9, fontweight="bold", pad=5)

    # Prefer non-review results for the scatter (richer signal)
    scatter_src = non_review_results or review_results
    if (scatter_src is not None
            and "df" in scatter_src
            and "Component_1" in scatter_src["df"].columns):
        sdf = scatter_src["df"]
        n_cl = int(sdf["Cluster"].nunique())
        cmap = matplotlib.colormaps["tab10"]
        ax_sc.scatter(sdf["Component_1"], sdf["Component_2"],
                      c=[cmap(int(c) % 10) for c in sdf["Cluster"]],
                      s=6, alpha=0.50, linewidths=0)
        # Annotate cluster centroids
        for cid in sorted(sdf["Cluster"].unique()):
            sub = sdf[sdf["Cluster"] == cid]
            ax_sc.text(sub["Component_1"].mean(), sub["Component_2"].mean(),
                       str(cid), fontsize=6, ha="center", va="center",
                       fontweight="bold",
                       color="white",
                       bbox=dict(boxstyle="round,pad=0.15",
                                 fc=cmap(int(cid) % 10), ec="none", alpha=0.8))
    else:
        ax_sc.text(0.5, 0.5, "Cluster data unavailable\n(run with clustering deps installed)",
                   transform=ax_sc.transAxes, ha="center", va="center",
                   fontsize=8, color="gray", style="italic")

    ax_sc.set_xlabel("PC 1", fontsize=7.5)
    ax_sc.set_ylabel("PC 2", fontsize=7.5)
    ax_sc.tick_params(labelsize=6)
    ax_sc.spines[["top", "right"]].set_visible(False)

    # -- 5  CLUSTER THEMES TABLE (right) --------------------------------------
    ax_th = fig.add_subplot(gs[2, 1])
    ax_th.set_axis_off()
    ax_th.set_title("Identified Research Themes", fontsize=9, fontweight="bold", pad=5)

    summaries = None
    if non_review_results and "summaries" in non_review_results:
        summaries = non_review_results["summaries"]
    elif review_results and "summaries" in review_results:
        summaries = review_results["summaries"]

    if summaries:
        sorted_cl = sorted(summaries.items(),
                           key=lambda x: x[1].get("size", 0), reverse=True)[:9]
        cmap2 = matplotlib.colormaps["tab10"]
        y = 0.97
        for cid, info in sorted_cl:
            terms = ", ".join(info.get("top_terms", [])[:6])
            size  = info.get("size", 0)
            col   = cmap2(int(cid) % 10)

            # Colour swatch
            ax_th.add_patch(mpatches.FancyBboxPatch(
                (0.00, y - 0.045), 0.025, 0.045,
                boxstyle="square,pad=0", fc=col, ec="none",
                transform=ax_th.transAxes,
            ))
            ax_th.text(0.035, y - 0.005,
                       f"C{cid}  (n={size:,})",
                       transform=ax_th.transAxes, ha="left", va="top",
                       fontsize=7, fontweight="bold", color="#1a3a5c")
            wrapped_terms = textwrap.fill(terms, 44)
            ax_th.text(0.035, y - 0.040,
                       wrapped_terms,
                       transform=ax_th.transAxes, ha="left", va="top",
                       fontsize=6, color="#444444", style="italic")
            y -= 0.115
            if y < 0.02:
                break
    else:
        ax_th.text(0.5, 0.5, "Clustering not available",
                   transform=ax_th.transAxes, ha="center", va="center",
                   fontsize=8, color="gray", style="italic")

    # -- 6  GAP-ANALYSIS BANNER ------------------------------------------------
    ax_g = fig.add_subplot(gs[3, :])
    ax_g.set_axis_off()

    has_gaps  = gap_indices is not None and len(gap_indices) > 0
    gap_bg    = "#fdf2f2" if has_gaps else "#f0fff4"
    gap_edge  = PALETTE["gap_red"]  if has_gaps else PALETTE["gap_green"]
    gap_title_col = PALETTE["gap_red"] if has_gaps else PALETTE["gap_green"]

    ax_g.add_patch(mpatches.FancyBboxPatch(
        (0, 0), 1, 1,
        boxstyle="round,pad=0.02", lw=1.5,
        ec=gap_edge, fc=gap_bg,
        transform=ax_g.transAxes,
    ))

    ax_g.text(0.5, 0.92,
              "Gap Analysis: Review Coverage of Research Topics",
              transform=ax_g.transAxes, ha="center", va="top",
              fontsize=9, fontweight="bold", color=gap_title_col)

    if gap_indices is None:
        body = ("Gap analysis requires the Review vs Non-Review comparison pipeline. "
                "Re-run with clustering dependencies installed.")
        ax_g.text(0.5, 0.50, body,
                  transform=ax_g.transAxes, ha="center", va="center",
                  fontsize=7.5, color="gray", style="italic")

    elif not has_gaps:
        ax_g.text(0.5, 0.50,
                  "No significant gaps detected -- existing reviews broadly cover all research clusters.",
                  transform=ax_g.transAxes, ha="center", va="center",
                  fontsize=8, color=PALETTE["gap_green"])

    else:
        n_total_clusters = (len(non_review_results["summaries"])
                            if non_review_results and "summaries" in non_review_results
                            else "?")
        pct_gap = (len(gap_indices) / n_total_clusters * 100
                   if isinstance(n_total_clusters, int) and n_total_clusters > 0 else 0)

        header_line = (f"{len(gap_indices)} gap cluster(s) of {n_total_clusters} "
                       f"({pct_gap:.0f}%) have cosine similarity < 0.30 to any review cluster:")
        ax_g.text(0.5, 0.82, header_line,
                  transform=ax_g.transAxes, ha="center", va="top",
                  fontsize=7.5, color="#333333")

        gap_bullets = []
        for idx in gap_indices[:4]:
            if summaries and idx in summaries:
                info  = summaries[idx]
                terms = ", ".join(info.get("top_terms", [])[:4])
                gap_bullets.append(
                    f"  C{idx} (n={info.get('size',0):,}):  {terms}"
                )

        # Draw bullet lines
        y = 0.62
        for bullet in gap_bullets:
            ax_g.text(0.03, y, bullet,
                      transform=ax_g.transAxes, ha="left", va="top",
                      fontsize=7, color="#333333")
            y -= 0.18

        if len(gap_indices) > 4:
            ax_g.text(0.03, y,
                      f"  ... and {len(gap_indices) - 4} additional gap(s).",
                      transform=ax_g.transAxes, ha="left", va="top",
                      fontsize=7, color="#555555", style="italic")

    # -- FOOTER ----------------------------------------------------------------
    fig.text(
        0.5, 0.005,
        (f"MedScrape | AI-Assisted Systematic Literature Review  |  "
         f"USMA Department of Mathematical Sciences  |  "
         f"Generated: {datetime.now().strftime('%Y-%m-%d')}"),
        ha="center", va="bottom",
        fontsize=5.5, color="#888888", style="italic",
    )

    # -- SAVE ------------------------------------------------------------------
    if output_path is None:
        output_path = f"vignette_search{search['id']}.pdf"
    plt.savefig(output_path, dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    print(f"    Saved: {os.path.basename(output_path)}")
    return output_path


# -----------------------------------------------------------------------------
#  DOCX Vignette — standalone figure objects + real Word text
# -----------------------------------------------------------------------------

def _chart_bytes(fig) -> io.BytesIO:
    """Save a matplotlib figure to a PNG BytesIO buffer and return it."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _fig_article_types(df_all: pd.DataFrame):
    """Horizontal bar chart of article types — returns (fig, ax)."""
    type_labels = ["Reviews", "Syst. Reviews", "Clin. Trials",
                   "Meta-Analyses", "RCTs", "Longitudinal"]
    type_cols   = ["Review", "SystematicReview", "ClinicalTrial",
                   "MetaAnalysis", "RCT", "LongitudinalStudy"]
    counts = [int(df_all[c].sum()) if c in df_all.columns else 0 for c in type_cols]

    fig, ax = plt.subplots(figsize=(3.8, 2.6))
    bars = ax.barh(type_labels, counts, color=PALETTE["type_colors"],
                   edgecolor="white", height=0.62)
    for bar, cnt in zip(bars, counts):
        if cnt > 0:
            ax.text(bar.get_width() + max(counts, default=1) * 0.03,
                    bar.get_y() + bar.get_height() / 2,
                    f"{cnt:,}", va="center", ha="left", fontsize=7)
    ax.set_xlabel("Count", fontsize=8)
    ax.set_title("Article-Type Distribution", fontsize=9, fontweight="bold")
    ax.tick_params(axis="y", labelsize=7.5)
    ax.tick_params(axis="x", labelsize=7)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_xlim(0, max(counts, default=1) * 1.22)
    fig.tight_layout()
    return fig


def _fig_trend(df_all: pd.DataFrame):
    """Publication trend bar + trend line — returns fig."""
    df_c = df_all[df_all["Year"].notna()].copy()
    fig, ax = plt.subplots(figsize=(3.8, 2.6))

    if not df_c.empty and df_c["Year"].nunique() >= 2:
        yc     = df_c["Year"].value_counts().sort_index()
        years  = yc.index.astype(int).values
        counts = yc.values
        ax.bar(years, counts, color=PALETTE["bar"], alpha=0.80, width=0.85)
        if len(years) >= 3:
            z  = np.polyfit(years, counts, 1)
            xs = np.linspace(years[0], years[-1], 200)
            ax.plot(xs, np.poly1d(z)(xs), "--",
                    color=PALETTE["trend"], lw=1.5, label="Trend")
            ax.legend(fontsize=6.5, loc="upper left", framealpha=0.7)
        if len(counts) >= 4:
            thirds = max(1, len(counts) // 3)
            early, late = counts[:thirds].mean(), counts[-thirds:].mean()
            if early > 0:
                pct = (late - early) / early * 100
                sym = "▲" if pct > 0 else "▼"
                col = PALETTE["gap_green"] if pct > 0 else PALETTE["gap_red"]
                ax.text(0.97, 0.95, f"{sym} {abs(pct):.0f}% growth",
                        transform=ax.transAxes, ha="right", va="top",
                        fontsize=7, color=col,
                        bbox=dict(boxstyle="round,pad=0.25", fc="white",
                                  ec="lightgray", alpha=0.85))
        tick_step = max(1, (years[-1] - years[0]) // 6)
        ticks = list(range(int(years[0]), int(years[-1]) + 1, tick_step))
        if int(years[-1]) not in ticks:
            ticks.append(int(years[-1]))
        ax.set_xticks(ticks)
        ax.set_xticklabels(ticks, rotation=45, ha="right", fontsize=6.5)
    else:
        ax.text(0.5, 0.5, "Insufficient year data",
                transform=ax.transAxes, ha="center", fontsize=9, color="gray")

    ax.set_ylabel("Publications / year", fontsize=8)
    ax.set_title("Publication Trend (2000-2025)", fontsize=9, fontweight="bold")
    ax.tick_params(axis="y", labelsize=6.5)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    return fig


def _fig_scatter(non_review_results, review_results):
    """PCA cluster scatter — returns fig."""
    fig, ax = plt.subplots(figsize=(3.8, 3.2))
    src = non_review_results or review_results
    if (src and "df" in src and "Component_1" in src["df"].columns):
        sdf  = src["df"]
        cmap = matplotlib.colormaps["tab10"]
        ax.scatter(sdf["Component_1"], sdf["Component_2"],
                   c=[cmap(int(c) % 10) for c in sdf["Cluster"]],
                   s=6, alpha=0.50, linewidths=0)
        for cid in sorted(sdf["Cluster"].unique()):
            sub = sdf[sdf["Cluster"] == cid]
            ax.text(sub["Component_1"].mean(), sub["Component_2"].mean(),
                    str(cid), fontsize=6, ha="center", va="center",
                    fontweight="bold", color="white",
                    bbox=dict(boxstyle="round,pad=0.15",
                              fc=cmap(int(cid) % 10), ec="none", alpha=0.85))
    else:
        ax.text(0.5, 0.5, "Cluster data unavailable",
                transform=ax.transAxes, ha="center", fontsize=9, color="gray",
                style="italic")
    ax.set_xlabel("PC 1", fontsize=8)
    ax.set_ylabel("PC 2", fontsize=8)
    ax.set_title("Research Cluster Map (PCA)", fontsize=9, fontweight="bold")
    ax.tick_params(labelsize=6)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    return fig


def generate_docx_vignette(search: dict,
                            df_all: pd.DataFrame,
                            review_results=None,
                            non_review_results=None,
                            similarity_matrix=None,
                            gap_indices=None,
                            output_path: str = None) -> str:
    """
    Build a structured DOCX where every chart is its own standalone picture
    object and all text (title, stats, themes, gaps) is real Word text.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io

    def _set_cell_bg(cell, hex_color: str):
        """Set table cell background colour."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.lstrip("#"))
        tcPr.append(shd)

    def _rgb(hex_color: str) -> RGBColor:
        h = hex_color.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _add_chart(cell, fig, width_in: float):
        """Embed a matplotlib figure as a picture inside a table cell."""
        buf  = _chart_bytes(fig)
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run  = para.add_run()
        run.add_picture(buf, width=Inches(width_in))

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width        = Inches(8.5)
    sec.page_height       = Inches(11)
    sec.left_margin       = Inches(0.85)
    sec.right_margin      = Inches(0.85)
    sec.top_margin        = Inches(0.75)
    sec.bottom_margin     = Inches(0.60)

    # ── Title ──────────────────────────────────────────────────────────────────
    h = doc.add_heading("", level=1)
    run = h.add_run(f"Search {search['id']}:  {search['title']}")
    run.font.color.rgb = _rgb(PALETTE["header_bg"])
    run.font.size      = Pt(16)
    h.alignment        = WD_ALIGN_PARAGRAPH.CENTER

    # Query
    qp = doc.add_paragraph()
    qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr = qp.add_run(f"PubMed Query:  {search['query']}")
    qr.italic         = True
    qr.font.size      = Pt(8.5)
    qr.font.color.rgb = _rgb("#444444")
    qp.paragraph_format.space_after = Pt(2)

    # Description
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(search["description"])
    dr.font.size      = Pt(8.5)
    dr.font.color.rgb = _rgb("#555555")
    dp.paragraph_format.space_after = Pt(6)

    # ── Stats row (4-cell shaded table) ───────────────────────────────────────
    df_c    = df_all[df_all["Year"].notna()]
    yr_min  = int(df_c["Year"].min()) if not df_c.empty else "—"
    yr_max  = int(df_c["Year"].max()) if not df_c.empty else "—"
    n_rev   = int(df_all.get("Review", pd.Series(dtype=int)).sum())

    stats = [
        ("Total Articles",  f"{len(df_all):,}"),
        ("Year Range",      f"{yr_min}–{yr_max}"),
        ("Unique Journals", f"{df_all['Journal'].nunique():,}"),
        ("Review Papers",   f"{n_rev:,}"),
    ]

    st = doc.add_table(rows=2, cols=4)
    st.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_i, (label, value) in enumerate(stats):
        # Value cell (row 0)
        vc = st.rows[0].cells[col_i]
        _set_cell_bg(vc, PALETTE["header_bg"])
        vp = vc.add_paragraph(value)
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.runs[0]
        vr.font.size      = Pt(16)
        vr.font.bold      = True
        vr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Label cell (row 1)
        lc = st.rows[1].cells[col_i]
        _set_cell_bg(lc, "2c5f8a")
        lp = lc.add_paragraph(label)
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.runs[0]
        lr.font.size      = Pt(8)
        lr.font.color.rgb = _rgb("#c8ddf0")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Charts row: Article Types | Publication Trend ─────────────────────────
    _h2 = doc.add_heading("Publication Overview", level=2)
    _h2.runs[0].font.color.rgb = _rgb(PALETTE["header_bg"])

    ct = doc.add_table(rows=1, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_chart(ct.rows[0].cells[0], _fig_article_types(df_all), 3.15)
    _add_chart(ct.rows[0].cells[1], _fig_trend(df_all),         3.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Cluster row: Scatter | Themes text ────────────────────────────────────
    _h3 = doc.add_heading("Research Cluster Analysis", level=2)
    _h3.runs[0].font.color.rgb = _rgb(PALETTE["header_bg"])

    cl = doc.add_table(rows=1, cols=2)
    cl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left — scatter plot
    _add_chart(cl.rows[0].cells[0], _fig_scatter(non_review_results, review_results), 3.15)

    # Right — cluster themes as real text
    right_cell = cl.rows[0].cells[1]
    summaries  = None
    if non_review_results and "summaries" in non_review_results:
        summaries = non_review_results["summaries"]
    elif review_results and "summaries" in review_results:
        summaries = review_results["summaries"]

    th_hdr = right_cell.add_paragraph("Identified Research Themes")
    th_hdr.runs[0].font.bold      = True
    th_hdr.runs[0].font.size      = Pt(9)
    th_hdr.runs[0].font.color.rgb = _rgb(PALETTE["header_bg"])
    th_hdr.paragraph_format.space_after = Pt(3)

    if summaries:
        cmap_t = matplotlib.colormaps["tab10"]
        for cid, info in sorted(summaries.items(),
                                 key=lambda x: x[1].get("size", 0), reverse=True)[:8]:
            terms  = ", ".join(info.get("top_terms", [])[:6])
            size   = info.get("size", 0)
            # Cluster label
            lp = right_cell.add_paragraph()
            lp.paragraph_format.space_before = Pt(3)
            lp.paragraph_format.space_after  = Pt(0)
            lr = lp.add_run(f"Cluster {cid}  (n={size:,})")
            lr.font.bold = True
            lr.font.size = Pt(8)
            rgba = cmap_t(int(cid) % 10)
            lr.font.color.rgb = RGBColor(
                int(rgba[0]*255), int(rgba[1]*255), int(rgba[2]*255))
            # Terms
            tp = right_cell.add_paragraph()
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after  = Pt(2)
            tr = tp.add_run(terms)
            tr.italic     = True
            tr.font.size  = Pt(7.5)
            tr.font.color.rgb = _rgb("#444444")
    else:
        right_cell.add_paragraph("Clustering data unavailable.")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Gap Analysis ───────────────────────────────────────────────────────────
    _h4 = doc.add_heading("Gap Analysis: Review Coverage of Research Topics", level=2)
    _h4.runs[0].font.color.rgb = _rgb(PALETTE["header_bg"])

    has_gaps = gap_indices is not None and len(gap_indices) > 0

    # Shaded banner paragraph
    gap_tbl = doc.add_table(rows=1, cols=1)
    gap_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    gap_cell = gap_tbl.rows[0].cells[0]
    _set_cell_bg(gap_cell, "fdf2f2" if has_gaps else "f0fff4")

    if gap_indices is None:
        gp = gap_cell.add_paragraph(
            "Gap analysis requires the Review vs Non-Review comparison pipeline.")
        gp.runs[0].italic     = True
        gp.runs[0].font.color.rgb = _rgb("#888888")

    elif not has_gaps:
        gp = gap_cell.add_paragraph(
            "No significant gaps detected — existing reviews broadly cover all research clusters.")
        gp.runs[0].font.bold      = True
        gp.runs[0].font.color.rgb = _rgb(PALETTE["gap_green"])

    else:
        n_total = (len(non_review_results["summaries"])
                   if non_review_results and "summaries" in non_review_results else "?")
        pct = (len(gap_indices) / n_total * 100
               if isinstance(n_total, int) and n_total > 0 else 0)

        hdr_p = gap_cell.add_paragraph()
        hdr_r = hdr_p.add_run(
            f"{len(gap_indices)} of {n_total} research clusters ({pct:.0f}%) "
            f"have cosine similarity < 0.30 to any review cluster:")
        hdr_r.font.bold      = True
        hdr_r.font.color.rgb = _rgb(PALETTE["gap_red"])
        hdr_r.font.size      = Pt(9)

        for idx in gap_indices[:5]:
            if summaries and idx in summaries:
                info  = summaries[idx]
                terms = ", ".join(info.get("top_terms", [])[:5])
                bp    = gap_cell.add_paragraph(style="List Bullet")
                br    = bp.add_run(f"Cluster {idx} (n={info.get('size',0):,}):  {terms}")
                br.font.size      = Pt(8.5)
                br.font.color.rgb = _rgb("#333333")
        if len(gap_indices) > 5:
            ep = gap_cell.add_paragraph()
            er = ep.add_run(f"… and {len(gap_indices) - 5} additional gap(s).")
            er.italic     = True
            er.font.size  = Pt(8)
            er.font.color.rgb = _rgb("#666666")

    # ── Footer ─────────────────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(8)
    fr = fp.add_run(
        f"MedScrape  |  AI-Assisted Systematic Literature Review  |  "
        f"USMA Department of Mathematical Sciences  |  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d')}"
    )
    fr.font.size      = Pt(7)
    fr.italic         = True
    fr.font.color.rgb = _rgb("#888888")

    # ── Save ───────────────────────────────────────────────────────────────────
    if output_path is None:
        output_path = f"vignette_search{search['id']}.docx"
    doc.save(output_path)
    print(f"    Saved: {os.path.basename(output_path)}")
    return output_path


# -----------------------------------------------------------------------------
#  Per-Search Pipeline
# -----------------------------------------------------------------------------

def run_search(search: dict,
               email: str,
               max_results: int,
               start_year: int,
               end_year: int,
               n_clusters: int,
               output_dir: str,
               data_dir: str) -> str | None:

    print(f"\n{'='*65}")
    print(f"  Search {search['id']}: {search['title']}")
    print(f"  Query : {search['query']}")
    print(f"{'='*65}")

    # Helper: load from cache CSV if it exists, otherwise scrape and save
    def _load_or_scrape(suffix: str, pub_type: str, label: str) -> pd.DataFrame:
        cache = os.path.join(data_dir, f"search{search['id']}_{search['short']}_{suffix}.csv")
        if os.path.exists(cache):
            df = pd.read_csv(cache)
            print(f"    Loaded from cache ({len(df):,} records): {os.path.basename(cache)}")
            return df
        df = _scrape(search["query"], email, max_results, start_year, end_year, pub_type)
        if not df.empty:
            df.to_csv(cache, index=False)
            print(f"    Saved: {os.path.basename(cache)}")
        return df

    # -- Scrape / load ALL papers ----------------------------------------------
    print(f"\n  [1/4]  Fetching ALL papers ...")
    df_all = _load_or_scrape("all", "all", "all")

    if df_all.empty:
        print("  No data retrieved. Skipping this search.")
        return None

    # -- Scrape / load REVIEWS and NON-REVIEWS for gap analysis ---------------
    print(f"\n  [2/4]  Fetching REVIEW papers ...")
    df_rev = _load_or_scrape("reviews", "review", "review")

    print(f"\n  [3/4]  Fetching NON-REVIEW (research) papers ...")
    df_nr = _load_or_scrape("nonreviews", "non-review", "non-review")

    # -- Clustering + Gap Analysis ---------------------------------------------
    review_results = non_review_results = similarity_matrix = gap_indices = None

    print(f"\n  [4/4]  Clustering & gap analysis ...")
    if CLUSTERING_OK:
        MIN_DOCS = max(n_clusters + 1, 10)

        # Cluster review papers
        if len(df_rev) >= MIN_DOCS:
            k_rev = min(n_clusters, max(2, len(df_rev) // 10))
            try:
                rc = EnhancedClusterer(df_rev.copy(), "Abstract")
                review_results = rc.cluster_and_reduce(
                    method="kmeans", n_clusters=k_rev,
                    embedding_type="tfidf", reduction_method="pca",
                )
                print(f"    Review clusters    : {review_results['n_clusters']}")
            except Exception as e:
                print(f"    [WARN] Review clustering failed: {e}")

        # Cluster non-review (primary research) papers
        if len(df_nr) >= MIN_DOCS:
            k_nr = min(n_clusters, max(2, len(df_nr) // 10))
            try:
                nrc = EnhancedClusterer(df_nr.copy(), "Abstract")
                non_review_results = nrc.cluster_and_reduce(
                    method="kmeans", n_clusters=k_nr,
                    embedding_type="tfidf", reduction_method="pca",
                )
                print(f"    Non-review clusters: {non_review_results['n_clusters']}")
            except Exception as e:
                print(f"    [WARN] Non-review clustering failed: {e}")

        # Gap analysis
        if review_results is not None and non_review_results is not None:
            try:
                comp = ReviewComparator(df_rev.copy(), df_nr.copy(), "Abstract")
                similarity_matrix = comp.compute_cosine_similarity_matrix(
                    review_results, non_review_results
                )
                gap_indices = comp.identify_gaps(similarity_matrix, threshold=0.3)
                print(f"    Research gaps      : {len(gap_indices)}")
            except Exception as e:
                print(f"    [WARN] Gap analysis failed: {e}")
    else:
        print("    Skipping (clustering_enhanced not available).")

    # -- Generate Vignettes (PDF + DOCX) ---------------------------------------
    base = f"vignette_search{search['id']}_{search['short']}"

    out_pdf = os.path.join(output_dir, f"{base}.pdf")
    draw_vignette(
        search=search,
        df_all=df_all,
        review_results=review_results,
        non_review_results=non_review_results,
        similarity_matrix=similarity_matrix,
        gap_indices=gap_indices,
        output_path=out_pdf,
    )

    out_docx = os.path.join(output_dir, f"{base}.docx")
    generate_docx_vignette(
        search=search,
        df_all=df_all,
        review_results=review_results,
        non_review_results=non_review_results,
        similarity_matrix=similarity_matrix,
        gap_indices=gap_indices,
        output_path=out_docx,
    )

    return out_pdf


# -----------------------------------------------------------------------------
#  Entry Point
# -----------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate MedScrape one-page vignettes for nutrition paper."
    )
    parser.add_argument("email", nargs="?",
                        help="NCBI-registered email (required by Entrez API)")
    parser.add_argument("--max",   type=int, default=30000,
                        help="Max PubMed results per query (default: 30000; no hard cap)")
    parser.add_argument("--start", type=int, default=2000,
                        help="Start year for date filter (default: 2000)")
    parser.add_argument("--end",   type=int, default=2025,
                        help="End year for date filter (default: 2025)")
    parser.add_argument("--clusters", type=int, default=8,
                        help="Max K-Means clusters (default: 8)")
    parser.add_argument("--search", type=int, default=0,
                        help="Run only this search ID 1-5 (default: 0 = all)")
    args = parser.parse_args()

    email = args.email or input("Enter your NCBI email address: ").strip()
    if not email or "@" not in email:
        print("Error: a valid email address is required for NCBI API access.")
        sys.exit(1)

    output_dir = os.path.join(_HERE, "vignettes")
    data_dir   = os.path.join(output_dir, "data")
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(data_dir, exist_ok=True)

    print("\n" + "="*65)
    print("  MedScrape Vignette Generator")
    print("="*65)
    print(f"  Email    : {email}")
    print(f"  Years    : {args.start}-{args.end}")
    print(f"  Max docs : {args.max:,} per query")
    print(f"  Clusters : up to {args.clusters}")
    print(f"  Output   : {output_dir}")
    print("="*65)

    searches = SEARCHES if args.search == 0 else [s for s in SEARCHES if s["id"] == args.search]
    if not searches:
        print(f"Error: search ID {args.search} not found (valid: 1-5).")
        sys.exit(1)

    generated = []
    for s in searches:
        try:
            path = run_search(
                search=s,
                email=email,
                max_results=args.max,
                start_year=args.start,
                end_year=args.end,
                n_clusters=args.clusters,
                output_dir=output_dir,
                data_dir=data_dir,
            )
            if path:
                generated.append(path)
        except KeyboardInterrupt:
            print("\n\nInterrupted by user.")
            break
        except Exception as e:
            print(f"\n  [ERROR] Search {s['id']} failed: {e}")
            import traceback
            traceback.print_exc()

    print(f"\n{'='*65}")
    print(f"  Done.  Generated {len(generated)} vignette(s):")
    for p in generated:
        print(f"    {p}")
    print("="*65)


if __name__ == "__main__":
    main()
